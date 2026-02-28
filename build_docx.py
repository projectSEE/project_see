"""
Build the System Flow and Implementation Details DOCX report.
Uses python-docx to create a professionally styled Word document
with embedded Mermaid diagram images.
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(BASE, "diagram_images")
OUT = os.path.join(BASE, "System_Flow_and_Implementation_Details.docx")

doc = Document()

# ─── Styles ──────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 5):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    if level == 1:
        hs.font.size = Pt(22)
        hs.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    elif level == 2:
        hs.font.size = Pt(16)
        hs.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    elif level == 3:
        hs.font.size = Pt(13)
    elif level == 4:
        hs.font.size = Pt(11)
        hs.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

# Page margins
for section in doc.sections:
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)


def add_para(text, bold=False, italic=False, size=None, color=None, align=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    if align: p.alignment = align
    if space_after is not None: p.paragraph_format.space_after = Pt(space_after)
    return p


def add_code(text, language=''):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    # Shading
    shading = run._element.get_or_add_rPr()
    sh = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:fill'): 'F1F5F9',
    })
    shading.append(sh)
    return p


def add_inline_code(paragraph, text):
    run = paragraph.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    return run


def add_image(filename, width=Inches(6.0)):
    path = os.path.join(IMG_DIR, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        add_para(f'[Diagram: {filename} not found]', italic=True, color=RGBColor(0x99, 0x33, 0x33))


def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9.5)
    # Rows
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)
    doc.add_paragraph()  # spacing


def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(10)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.2 * level)
    return p


# ═══════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
add_para('System Flow and\nImplementation Details',
         bold=True, size=28, color=RGBColor(0x25, 0x63, 0xEB),
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

add_para('SEE (Visual Assistant) — Flutter Mobile Application',
         size=12, color=RGBColor(0x64, 0x74, 0x8B),
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para('Code-level execution paths, data flows, and implementation mechanics',
         size=10, color=RGBColor(0x94, 0xA3, 0xB8),
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para('Codebase analysed: lib/ (35 Dart files), android/.../MainActivity.kt',
         size=10, color=RGBColor(0x94, 0xA3, 0xB8),
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════
# 1. SYSTEM FLOWCHARTS
# ═══════════════════════════════════════════════════════════

doc.add_heading('1. System Flowcharts', level=1)

# 1.1
doc.add_heading('1.1 Sensor & Hardware Integration — Camera Frame Pipeline', level=2)
add_para('The camera pipeline has two distinct consumers: ML Kit Object Detection (processed on every frame) and the Gemini Live API (sampled at 1 fps and converted to JPEG via a native platform channel).')
add_image('camera_pipeline.png')

# 1.2
doc.add_heading('1.2 Microphone Audio Capture Pipeline (16 kHz PCM with Hardware AEC)', level=2)
add_image('audio_pipeline.png')

# 1.3
doc.add_heading('1.3 On-Device ML Inference — Depth Estimation & Object Detection', level=2)
add_image('ml_inference.png')

# 1.4
doc.add_heading('1.4 Real-Time Bidirectional Streaming — Gemini Live API', level=2)
add_image('gemini_live.png')

doc.add_heading('Barge-In Handling Detail', level=4)
add_para('When the server signals msg.interrupted == true:')
add_bullet('AudioOutput.stopImmediately() is called, which invokes the native flush method.')
add_bullet('The native Kotlin code executes audioTrack.pause() → audioTrack.flush() → audioTrack.play(), instantly clearing the playback buffer.')
add_bullet('The _aiIsSpeaking flag is reset and _isFirstAudioChunk is set to true so the next response triggers a fresh earcon.')

# 1.5
doc.add_heading('1.5 Background Fall Detection and SOS Sequence', level=2)
add_image('fall_detection.png')

# 1.6
doc.add_heading('1.6 Chat Message Flow with POI Extraction', level=2)
add_image('chat_flow.png')


# ═══════════════════════════════════════════════════════════
# 2. IMPLEMENTATION DETAILS
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
doc.add_heading('2. Implementation Details', level=1)

# 2.1 State Management
doc.add_heading('2.1 State Management', level=2)
add_para('The application uses Flutter\'s built-in StatefulWidget/setState pattern exclusively. There is no provider, bloc, or riverpod dependency.')

add_table(
    ['Domain', 'State Container', 'Key Variables'],
    [
        ['Theme', 'ThemeNotifier (singleton ChangeNotifier)', '_themeMode (light/dark)'],
        ['Fall Detection', '_SafetyMonitorState', '_isFreeFalling, _freeFallTimestamp, _isAlertActive, _emergencyPhone'],
        ['Object Detection', '_ObjectDetectionScreenState', '_obstacles, _isDetecting, _depthModelLoaded, _streamErrorCount'],
        ['Live Session', '_LiveScreenState', '_isConnected, _isStreamingAudio/Video, _aiIsSpeaking, _pushToTalkMode, _pttPressed'],
        ['Chat', '_ChatScreenState', '_messages (list), _isLiveMode, _isRecording, _isAiSpeaking, _pendingImageBytes'],
        ['Depth Estimation', 'DepthEstimationService', '_depthHistory (Map<String, Queue>), _cachedResult, _isProcessing'],
        ['ML Kit', 'MLKitService', '_isProcessing (frame-level mutex), _frameCount'],
        ['TTS', 'TTSService (singleton)', '_isReady, _isSpeaking, _lastSpoken, _lastSpokenTime'],
    ]
)

add_para('High-frequency sensor streams (camera, accelerometer) are throttled explicitly:', bold=True)
add_bullet('Depth estimation: Minimum 333 ms between inferences (max ~3 FPS) via _minInterval.')
add_bullet('Vibration: 300 ms throttle via _lastVibration timestamp comparison.')
add_bullet('TTS obstacle announcements: 3-second Timer.periodic in ObjectDetectionScreen.')
add_bullet('Video frames to Gemini: 1-second Timer.periodic in LiveScreen.')
add_bullet('Location awareness POI announcements: 5-second cooldown in LocationAwarenessService.')

# 2.2 Data Structures
doc.add_heading('2.2 Data Structures & Transformation', level=2)

doc.add_heading('YUV420 → NV21 Conversion (Dart, for ML Kit)', level=3)
add_para('Located in ml_kit_service.dart. Copies Y plane row-by-row (respecting bytesPerRow stride), then interleaves V and U planes in VUVU ordering to produce NV21 format.')
add_code('''Uint8List _convertYUV420ToNV21(CameraImage image) {
  final int ySize = width * height;
  final Uint8List nv21 = Uint8List(ySize + uvSize);
  // Copy Y plane row-by-row
  for (int y = 0; y < height; y++)
    for (int x = 0; x < width; x++)
      nv21[y * width + x] = yPlane.bytes[y * yPlane.bytesPerRow + x];
  // Interleave V then U (NV21 = VUVU ordering)
  int uvIndex = ySize;
  for (int y = 0; y < uvHeight; y++)
    for (int x = 0; x < uvWidth; x++) {
      nv21[uvIndex++] = vPlane.bytes[vIdx]; // V first
      nv21[uvIndex++] = uPlane.bytes[uIdx]; // then U
    }
  return nv21;
}''')

doc.add_heading('YUV420 → JPEG (Native Kotlin, for Gemini video)', level=3)
add_para('Located in MainActivity.kt. Uses Android\'s YuvImage.compressToJpeg() after constructing an NV21 byte array from three YUV planes received via MethodChannel. Quality is set to 40 for bandwidth efficiency.')

doc.add_heading('YUV420 → RGB → Float32 NCHW (Dart, for ONNX Depth)', level=3)
add_para('Located in depth_estimation_service.dart. Performs nearest-neighbour resize to 252×252, normalises to [0.0, 1.0], and arranges in NCHW layout (channel-first).')
add_code('''Float32List _preprocessImageFloat32(Uint8List rgb, int w, int h) {
  final data = Float32List(1 * 3 * 252 * 252);
  for (int c = 0; c < 3; c++)
    for (int y = 0; y < 252; y++)
      for (int x = 0; x < 252; x++) {
        final srcX = (x * w / 252).round().clamp(0, w - 1);
        final srcY = (y * h / 252).round().clamp(0, h - 1);
        data[(c*252*252)+(y*252)+x] = rgb[(srcY*w+srcX)*3+c] / 255.0;
      }
  return data;
}''')

doc.add_heading('Temporal Depth Trend — Sliding Window', level=3)
add_bullet('Per-object label Queue<DepthSample> with a maximum of 5 samples.')
add_bullet('Trend: changeRate = (currentDepth − oldestDepth) / timeDelta.')
add_bullet('Thresholds: >0.15/s → approaching_fast; >0.05/s → approaching; <−0.15/s → moving_away_fast.')
add_bullet('isDanger flag: isApproaching && currentDepth > 0.5.')
add_bullet('Stale entries (>3 s since last sample) cleaned via _cleanupOldHistory().')

doc.add_heading('Firestore Data Models', level=3)
add_code('''conversations/{userId}/topics/{topicId}
    ├── createdAt, lastUpdated, firstMessage, lastMessage
    └── messages/{messageId}
            ├── role: 'user' | 'assistant' | 'system'
            ├── content, timestamp, imageUrl?, hasImage

users/{fullName}
    ├── profile: { fullName, email, phone, emergencyContact* }
    └── settings: { visualImpairment, hearingImpairment, ... }

pois/{poiName}
    ├── name, type, description, safetyNotes
    ├── location: { lat, lng, address }
    └── addedBy, createdAt

fall_detection_feedback/{docId}
    ├── userPhone, triggerCorrect, callCorrect, timestamp

traffic_feedback/{docId}
    ├── wasSuccessful, timestamp, deviceTime''')

doc.add_heading('POI Extraction from Gemini Response', level=3)
add_para('The chat screen parses Gemini responses for ADD_POI markers using brace-counting (not regex) to reliably extract nested JSON.')

# 2.3 Error Handling
doc.add_heading('2.3 Error Handling & Edge Cases', level=2)

doc.add_heading('Camera & Sensor Permissions', level=3)
add_table(
    ['Scenario', 'Handling'],
    [
        ['Camera permission denied', 'Full-screen card with "Open Settings" button via openAppSettings().'],
        ['Location permission denied', 'Geolocator.requestPermission(); if permanently denied, TTS + retry/settings dialog.'],
        ['Location permanently denied', 'AlertDialog with "Open Settings" option.'],
        ['Microphone permission denied', 'AudioInput throws Exception(\'Microphone permission not granted\').'],
    ]
)

doc.add_heading('Network & API Failures', level=3)
add_table(
    ['Scenario', 'Handling'],
    [
        ['Gemini API error (chat)', 'Returns "Error communicating with AI service" — displayed as model message.'],
        ['Live API WebSocket closed', '_receiveLoop() detects \'Closed\' and calls _stopAll() to tear down all streams.'],
        ['Custom VAD connect fails', 'Falls back to _sdkFallbackConnect() using standard Firebase AI SDK.'],
        ['Places API (New) HTTP error', 'Falls back to _searchPlacesOldApi() using legacy Nearby Search endpoint.'],
        ['Location timeout', '15-second timeout; falls back to Geolocator.getLastKnownPosition().'],
    ]
)

doc.add_heading('ML Inference Failures', level=3)
add_table(
    ['Scenario', 'Handling'],
    [
        ['Stream mode repeated failures', 'After 100 consecutive empty results, switches to file-based detection (2 s Timer).'],
        ['ONNX model load failure', 'Returns false; object detection continues without depth data.'],
        ['Depth estimation busy', 'Returns _cachedResult when within 333 ms throttle window.'],
    ]
)

doc.add_heading('TTS Engine Binding (Xiaomi/MIUI)', level=3)
add_para('TTSService.initialize() retries configuration up to 5 times with increasing delays (1 s, 2 s, 3 s, 4 s, 5 s). If all attempts fail, _isReady is still set to true so that speak() can retry at call time.')

doc.add_heading('App Lifecycle', level=3)
add_bullet('inactive: Calls _stopAll() / cancels timers and disposes camera.')
add_bullet('resumed: Re-initialises the camera.')

# 2.4 Package Implementation
doc.add_heading('2.4 Package Implementation', level=2)

doc.add_heading('record (Audio Capture)', level=3)
add_para('Configured in audio_input.dart:')
add_bullet('voiceCommunication enables Android\'s AcousticEchoCanceler, NoiseSuppressor, and AutomaticGainControl.')
add_bullet('modeInCommunication optimises the audio pipeline for two-way communication.')
add_bullet('muteAudio: false ensures AI speech continues to play through the speaker.')
add_code('''RecordConfig(
  encoder: AudioEncoder.pcm16bits,
  numChannels: 1, sampleRate: 16000,
  androidConfig: AndroidRecordConfig(
    audioSource: AndroidAudioSource.voiceCommunication,
    audioManagerMode: AudioManagerMode.modeInCommunication,
    muteAudio: false))''')

doc.add_heading('Native AudioTrack (AI Speech Playback)', level=3)
add_para('The AudioOutput class communicates with native Kotlin via MethodChannel(\'audio_output_channel\'). The native implementation in MainActivity.kt:')
add_bullet('Sets AudioManager.MODE_IN_COMMUNICATION globally.')
add_bullet('Creates AudioTrack with USAGE_VOICE_COMMUNICATION at 24 kHz (Gemini\'s output rate).')
add_bullet('Attaches AcousticEchoCanceler and NoiseSuppressor to the audio session.')
add_bullet('Forces routing to the built-in loudspeaker using setCommunicationDevice() on Android 12+.')
add_bullet('Maximises STREAM_VOICE_CALL volume for blind users.')
add_para('The flush method (barge-in): audioTrack.pause() → audioTrack.flush() → audioTrack.play().')

doc.add_heading('Firebase AI SDK (firebase_ai)', level=3)
add_bullet('Chat mode: generativeModel(\'gemini-2.5-flash\') with generateContent().')
add_bullet('Live mode: liveGenerativeModel(\'gemini-live-2.5-flash-native-audio\') with ResponseModalities.audio.')
add_bullet('Custom VAD: AccessibleLiveConnector bypasses SDK connect() with manual WebSocket + VAD config.')

doc.add_heading('flutter_tts (Text-to-Speech)', level=3)
add_bullet('Singleton (TTSService._instance). Language: en-US, rate: 0.5, volume: 1.0.')
add_bullet('Duplicate suppression: same message within 2 s is skipped (unless force: true).')
add_bullet('awaitSpeakCompletion(true) enables synchronous speech.')

doc.add_heading('onnxruntime (Depth Estimation)', level=3)
add_bullet('Model: assets/models/depth_anything_v2.onnx (Depth Anything V2, ViT-S).')
add_bullet('Input: Float32List [1, 3, 252, 252] (NCHW). 252 = 18 × 14 (ViT patch size).')
add_bullet('Inference: OrtSession.fromBuffer() → session.runAsync().')
add_bullet('Output: Depth map 252×252, normalised [0.0, 1.0] via min-max.')

doc.add_heading('google_mlkit_object_detection', level=3)
add_bullet('Stream mode, classifyObjects: true, multipleObjects: true.')
add_bullet('Position: left / center / right based on centerX vs. image width thirds.')
add_bullet('Proximity: relativeSize = objectArea / imageArea. isClose > 0.10; isVeryClose > 0.25.')

doc.add_heading('vibration + HapticFeedback', level=3)
add_table(
    ['Proximity', 'Method', 'Pattern'],
    [
        ['> 0.10 (very close)', '_vibrateHeavy()', 'heavyImpact × 2 (100 ms gap)'],
        ['> 0.05 (close)', '_vibrateMedium()', 'mediumImpact × 1'],
        ['> 0.01 (detected)', '_vibrateLight()', 'lightImpact × 1'],
    ]
)
add_para('Approaching objects receive an intensityBoost of 0.2 added to raw proximity, clamped to [0.0, 1.0].')

# 2.5 Earcon System
doc.add_heading('2.5 Earcon System (Accessibility Cues)', level=2)
add_table(
    ['Event', 'Earcon', 'Implementation'],
    [
        ['AI starts listening', '_playListeningEarcon()', 'HapticFeedback.lightImpact() + SystemSound.play(click)'],
        ['User releases PTT', '_playProcessingEarcon()', 'HapticFeedback.mediumImpact() × 2 (100 ms apart)'],
        ['AI starts responding', '_playResponseEarcon()', 'HapticFeedback.heavyImpact() + SystemSound.play(click)'],
    ]
)

# 2.6 Traffic Light Monitoring
doc.add_heading('2.6 Traffic Light Monitoring', level=2)
add_bullet('Sends initial prompt asking Gemini to identify pedestrian crossing signal colour.')
add_bullet('Re-prompts every 5 seconds via Timer.periodic; Gemini only speaks on colour change.')
add_bullet('On stop, collects user feedback and uploads to Firestore collection(\'traffic_feedback\').')

# 2.7 Authentication Flow
doc.add_heading('2.7 Authentication Flow', level=2)
add_image('auth_flow.png', width=Inches(5.0))

# 2.8 Navigation
doc.add_heading('2.8 Navigation & Location Awareness', level=2)
add_bullet('Route calculation: NavigationService.getRoute() calls Google Directions API (walking mode).')
add_bullet('Live tracking: Geolocator.getPositionStream(distanceFilter: 5). Advances step when within 10 m.')
add_bullet('Explore mode: LocationAwarenessService tracks with 10 m filter, calls Places API within 100 m radius. Announced POIs stored in Set<String> (max 50).')

# Footer
doc.add_paragraph()
add_para('Report generated from codebase analysis on 26 February 2026.',
         italic=True, size=9, color=RGBColor(0x94, 0xA3, 0xB8),
         align=WD_ALIGN_PARAGRAPH.CENTER)


# ─── Save ──────────────────────────────────────────────
doc.save(OUT)
print(f"DOCX saved to: {OUT}")
print(f"Size: {os.path.getsize(OUT) / 1024:.1f} KB")
